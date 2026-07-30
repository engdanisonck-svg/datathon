from __future__ import annotations

import json
import sys
from pathlib import Path

import pandas as pd
import pdfplumber
from docx import Document


def profile_xlsx(path: Path) -> dict:
    book = pd.ExcelFile(path)
    result = {"path": str(path), "sheets": {}}
    for sheet in book.sheet_names:
        frame = pd.read_excel(path, sheet_name=sheet)
        columns = []
        for name in frame.columns:
            series = frame[name]
            sample = series.dropna().astype(str).head(8).tolist()
            columns.append(
                {
                    "name": str(name),
                    "dtype": str(series.dtype),
                    "non_null": int(series.notna().sum()),
                    "null": int(series.isna().sum()),
                    "unique": int(series.nunique(dropna=True)),
                    "sample": sample,
                }
            )
        result["sheets"][sheet] = {
            "rows": int(len(frame)),
            "cols": int(len(frame.columns)),
            "duplicates": int(frame.duplicated().sum()),
            "columns": columns,
        }
    return result


def extract_pdf(path: Path) -> dict:
    pages = []
    with pdfplumber.open(path) as pdf:
        for number, page in enumerate(pdf.pages, 1):
            pages.append({"page": number, "text": page.extract_text() or ""})
    return {"path": str(path), "pages": pages}


def extract_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [
        {"style": p.style.name if p.style else "", "text": p.text}
        for p in doc.paragraphs
        if p.text.strip()
    ]
    tables = []
    for table in doc.tables:
        tables.append([[cell.text for cell in row.cells] for row in table.rows])
    return {"path": str(path), "paragraphs": paragraphs, "tables": tables}


def main() -> None:
    xlsx, pdf, docx, output = map(Path, sys.argv[1:5])
    report = {
        "xlsx": profile_xlsx(xlsx),
        "pdf": extract_pdf(pdf),
        "docx": extract_docx(docx),
    }
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
